"""
.md 编译稿 → .docx 排版输出

用法:
    python md_to_docx.py <input.md> [output_dir]

格式规范: FORMAT_SPEC.md（公文格式标准）
"""

import re
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ═══════════════════════════════════════════
#  公文格式常量（来源于 FORMAT_SPEC.md）
# ═══════════════════════════════════════════

# 页面
PAGE_WIDTH = Cm(21)
PAGE_HEIGHT = Cm(29.7)
MARGIN_TOP = Cm(2.54)
MARGIN_BOTTOM = Cm(2.54)
MARGIN_LEFT = Cm(3.17)
MARGIN_RIGHT = Cm(3.17)

# 字体
FONT_TITLE_CN = "方正小标宋简体"
FONT_H1_CN = "黑体"
FONT_H2_CN = "楷体_GB2312"
FONT_BODY_CN = "仿宋_GB2312"
FONT_CITATION_CN = "楷体_GB2312"
FONT_EN = "Times New Roman"

# 字号
SIZE_TITLE = Pt(22)      # 二号
SIZE_H1 = Pt(16)         # 三号
SIZE_H2 = Pt(16)         # 三号
SIZE_BODY = Pt(16)       # 三号
SIZE_CITATION = Pt(12)   # 小四
SIZE_TAG = Pt(9)         # 小五

# 段落
LINE_SPACING = Pt(28)        # 固定 28 磅
FIRST_LINE_INDENT = Cm(1.13)  # 三号仿宋 2 字符 ≈ 32pt ≈ 1.13cm


def _setup_page(doc):
    """页面设置"""
    sec = doc.sections[0]
    sec.page_width = PAGE_WIDTH
    sec.page_height = PAGE_HEIGHT
    sec.top_margin = MARGIN_TOP
    sec.bottom_margin = MARGIN_BOTTOM
    sec.left_margin = MARGIN_LEFT
    sec.right_margin = MARGIN_RIGHT


def _set_run(run, cn_font, font_size, bold=False, en_font=FONT_EN):
    """统一设置 run 的字体和字号"""
    run.font.size = font_size
    run.font.bold = bold
    run.font.name = en_font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:eastAsia'), cn_font)
    rPr.insert(0, rFonts)


def _add_paragraph(doc, text, cn_font, font_size, bold=False,
                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   first_line_indent=None, space_after=Pt(0)):
    """添加段落，返回 paragraph 对象"""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = space_after
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    run = p.add_run(text)
    _set_run(run, cn_font, font_size, bold)
    return p


def _add_title(doc, text):
    """报告标题：上下各空一行，居中，方正小标宋简体，二号"""
    _add_paragraph(doc, "", FONT_BODY_CN, SIZE_BODY)  # 上空一行
    p = _add_paragraph(doc, text, FONT_TITLE_CN, SIZE_TITLE,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "", FONT_BODY_CN, SIZE_BODY)  # 下空一行
    return p


def _add_citation(doc, text):
    """引用行：居中，楷体小四，灰色"""
    p = _add_paragraph(doc, text, FONT_CITATION_CN, SIZE_CITATION,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    return p


def _add_h1(doc, text):
    """一级标题：黑体三号，首行缩进 2 字符"""
    return _add_paragraph(doc, text, FONT_H1_CN, SIZE_H1,
                          first_line_indent=FIRST_LINE_INDENT)


def _add_h2(doc, text):
    """二级标题：楷体_GB2312 三号，首行缩进 2 字符"""
    return _add_paragraph(doc, text, FONT_H2_CN, SIZE_H2,
                          first_line_indent=FIRST_LINE_INDENT)


def _add_body(doc, text):
    """正文：仿宋三号，首行缩进 2 字符"""
    if not text.strip():
        return None
    return _add_paragraph(doc, text.strip(), FONT_BODY_CN, SIZE_BODY,
                          first_line_indent=FIRST_LINE_INDENT)


def _add_abstract(doc, text):
    """摘要标签+正文：标签加粗，正文正常"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    # 标签部分
    run_label = p.add_run("摘要：")
    _set_run(run_label, FONT_BODY_CN, SIZE_BODY, bold=True)
    # 正文部分
    run_body = p.add_run(text)
    _set_run(run_body, FONT_BODY_CN, SIZE_BODY, bold=False)
    return p


def _add_keywords(doc, text):
    """关键词标签+内容：标签加粗，内容正常"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    run_label = p.add_run("关键词：")
    _set_run(run_label, FONT_BODY_CN, SIZE_BODY, bold=True)
    run_body = p.add_run(text)
    _set_run(run_body, FONT_BODY_CN, SIZE_BODY, bold=False)
    return p


# ═══════════════════════════════════════════
#  Markdown 解析
# ═══════════════════════════════════════════

def parse_md(filepath: str) -> dict:
    """解析 .md 编译稿"""
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()

    result = {
        "title": "",
        "citation": "",
        "abstract": "",
        "keywords": "",
        "sections": [],  # [(level, title, body), ...]
        "tags": [],
    }

    # 跳过 YAML frontmatter
    body = content
    if content.startswith("---"):
        end = content.find("---", 3)
        if end != -1:
            body = content[end + 3:]

    lines = body.strip().split("\n")

    # 提取标题
    for i, line in enumerate(lines):
        s = line.strip()
        if s.startswith("# ") and not s.startswith("## "):
            result["title"] = s[2:].strip()
            if i + 1 < len(lines) and lines[i + 1].strip().startswith(">"):
                result["citation"] = lines[i + 1].strip().lstrip("> ").strip()
            break

    # 提取摘要和关键词
    for i, line in enumerate(lines):
        s = line.strip()
        if s.startswith("**摘要：**"):
            result["abstract"] = s.replace("**摘要：**", "").strip()
        elif s.startswith("**关键词：**"):
            result["keywords"] = s.replace("**关键词：**", "").strip()

    # 按 ## 分割章节
    sections_raw = re.split(r'\n(?=## )', body)

    for section in sections_raw:
        section = section.strip()
        header_match = re.match(r'##\s+(.+)', section)
        if not header_match:
            continue
        header = header_match.group(1).strip()
        section_body = section[header_match.end():].strip()

        if "标签" in header:
            result["tags"] = re.findall(r'#([\w一-鿿]+)', section_body)
            continue

        # 判断层级
        if re.match(r'[一二三四五六七八九十]、', header):
            level = 1
        elif re.match(r'（[一二三四五六七八九十]）', header):
            level = 2
        elif re.match(r'\d+\.', header):
            level = 2
        else:
            level = 1

        result["sections"].append((level, header, section_body))

    return result


# ═══════════════════════════════════════════
#  .docx 生成
# ═══════════════════════════════════════════

def generate_report(data: dict, output_path: str) -> None:
    """生成公文格式分析报告 .docx"""
    doc = Document()
    _setup_page(doc)

    # 标题
    if data.get("title"):
        _add_title(doc, data["title"])

    # 引用行
    if data.get("citation"):
        _add_citation(doc, data["citation"])

    # 摘要
    abstract = data.get("abstract", "")
    if abstract:
        _add_abstract(doc, abstract)

    # 关键词
    keywords = data.get("keywords", "")
    if keywords:
        _add_keywords(doc, keywords)

    _add_paragraph(doc, "", FONT_BODY_CN, SIZE_BODY)

    # 章节
    sections = data.get("sections", [])
    for level, title, body_text in sections:
        if level == 1:
            _add_h1(doc, title)
        elif level == 2:
            _add_h2(doc, title)

        if body_text:
            for para in body_text.split("\n"):
                stripped = para.strip()
                if stripped and not stripped.startswith("（") and not re.match(r'[一二三四五六七八九十]、', stripped):
                    _add_body(doc, stripped)
                elif stripped:
                    # 子标题行已在上面处理
                    pass

            # 子节内容按 (一) (二) 拆分并独立处理
            if level == 1:
                sub_parts = re.split(r'\n(?=（[一二三四五六七八九十]）)', body_text)
                for sp in sub_parts:
                    sp = sp.strip()
                    if not sp:
                        continue
                    h2_match = re.match(r'（([一二三四五六七八九十])）(.+)', sp)
                    if h2_match:
                        sub_num = h2_match.group(1)
                        rest_text = h2_match.group(2)
                        sub_title = f"（{sub_num}）{rest_text.split(chr(10))[0].strip()}"
                        _add_h2(doc, sub_title)
                        rest = "\n".join(rest_text.split("\n")[1:]).strip()
                        for rp in rest.split("\n"):
                            rp = rp.strip()
                            if rp and not re.match(r'（[一二三四五六七八九十]）', rp):
                                _add_body(doc, rp)
                    else:
                        for rp in sp.split("\n"):
                            rp = rp.strip()
                            if rp:
                                _add_body(doc, rp)

    # 标签
    tags = data.get("tags", [])
    if tags:
        _add_paragraph(doc, "", FONT_BODY_CN, SIZE_BODY)
        tag_text = "标签: " + " | ".join(f"#{t}" for t in tags)
        p = _add_paragraph(doc, tag_text, FONT_BODY_CN, SIZE_TAG,
                           first_line_indent=Cm(0))
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)


# ═══════════════════════════════════════════
#  主入口
# ═══════════════════════════════════════════

def convert(md_path: str, output_dir: str = "") -> str:
    """将 .md 编译稿转换为 .docx 报告。同名文件自动递增版本号。"""
    if not os.path.exists(md_path):
        raise FileNotFoundError(f"文件不存在: {md_path}")

    if not output_dir:
        output_dir = os.path.dirname(md_path)
    os.makedirs(output_dir, exist_ok=True)

    data = parse_md(md_path)
    basename = os.path.splitext(os.path.basename(md_path))[0]

    # 同名文件检测：已存在则加 (1)、(2)...
    output_path = os.path.join(output_dir, f"{basename}.docx")
    counter = 1
    while os.path.exists(output_path):
        output_path = os.path.join(output_dir, f"{basename}（{counter}）.docx")
        counter += 1

    generate_report(data, output_path)

    return output_path


if __name__ == "__main__":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

    if len(sys.argv) < 2:
        print("用法: python md_to_docx.py <input.md> [output_dir]")
        sys.exit(1)

    md_path = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else ""
    result = convert(md_path, output_dir)
    print(f"已导出: {result}")
